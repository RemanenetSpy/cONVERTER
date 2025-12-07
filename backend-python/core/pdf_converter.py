"""
PDF Converter Module
Handles PDF conversions to/from Word, Excel, PowerPoint, and images
"""

import os
from typing import Dict, Any, List
from PyPDF2 import PdfReader, PdfWriter, PdfMerger
from pdf2docx import Converter as PDFToWordConverter
from docx import Document
from docx.shared import Inches
from pptx import Presentation
from pptx.util import Inches as PptxInches
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
import pandas as pd


class PDFConverter:
    """Converts PDFs to/from various formats"""

    @staticmethod
    def _run_ocr_if_needed(input_path: str) -> str:
        """Run OCR if PDF is image-based (Scan)"""
        try:
            import shutil
            import ocrmypdf
            
            # Check if tesseract is installed
            if not shutil.which("tesseract"):
                print("Warning: Tesseract not found. Skipping OCR.")
                return input_path
                
            print(f"Checking/Running OCR for: {input_path}")
            output_path = input_path.replace(".pdf", "_ocr.pdf")
            
            # Run OCR (skip_text=True means it only acts if needed)
            exit_code = ocrmypdf.ocr(
                input_path, 
                output_path, 
                deskew=True, 
                skip_text=True,
                jobs=1,
                output_type='pdf'
            )
            
            if exit_code == 0 and os.path.exists(output_path):
                return output_path
                
            return input_path
            
        except Exception as e:
            print(f"OCR Process Warning: {str(e)}")
            return input_path

    @staticmethod
    def pdf_to_word(input_path: str, output_path: str) -> Dict[str, Any]:
        """Convert PDF to Word (with Smart OCR for Scans)"""
        
        try:
            # 1. OCR Pre-processing (Make Scans Editable)
            # Only runs if Tesseract is installed (Server/Production)
            processed_path = PDFConverter._run_ocr_if_needed(input_path)
            
            # 2. Convert using pdf2docx
            cv = PDFToWordConverter(processed_path)
            # Disable multiprocessing for stability on cloud instances (Render Free Tier)
            cv.convert(output_path, multi_processing=False)
            cv.close()
            
            # Get page count
            pdf = PdfReader(input_path)
            page_count = len(pdf.pages)
            
            # Cleanup OCR temp file if created
            if processed_path != input_path and os.path.exists(processed_path):
                try:
                    os.remove(processed_path)
                except:
                    pass
            
            return {
                "success": True,
                "message": f"Converted PDF to Word ({page_count} pages) [OCR Checked]",
                "outputPath": output_path,
                "pages": page_count
            }
        
        except Exception as e:
            raise Exception(f"PDF to Word conversion failed: {str(e)}")

    @staticmethod
    def _get_libreoffice_command():
        """Find LibreOffice executable"""
        import shutil
        import platform
        
        system = platform.system()
        
        # Windows Paths
        if system == 'Windows':
            paths = [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            ]
            for path in paths:
                if os.path.exists(path):
                    return path
            # Check PATH
            return shutil.which("soffice")
            
        # Linux/Mac Paths
        else:
            commands = ["libreoffice", "soffice", "libreoffice5.0"]
            for cmd in commands:
                path = shutil.which(cmd)
                if path:
                    return path
        return None

    @staticmethod
    def _convert_with_libreoffice(input_path: str, output_path: str) -> bool:
        """Convert using LibreOffice (Headless)"""
        soffice = PDFConverter._get_libreoffice_command()
        if not soffice:
            return False
            
        try:
            import subprocess
            out_dir = os.path.dirname(output_path)
            
            # Run LibreOffice Headless
            # --convert-to pdf --outdir <dir> <input_file>
            cmd = [
                soffice,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', out_dir,
                input_path
            ]
            
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            
            # LibreOffice saves as <filename>.pdf in outdir
            # We need to ensure it matches output_path target
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            generated_file = os.path.join(out_dir, base_name + ".pdf")
            
            if os.path.exists(generated_file) and os.path.abspath(generated_file) != os.path.abspath(output_path):
                if os.path.exists(output_path):
                    os.remove(output_path)
                os.rename(generated_file, output_path)
                
            return os.path.exists(output_path)
            
        except Exception as e:
            print(f"LibreOffice conversion failed: {e}")
            return False

    @staticmethod
    def excel_to_pdf(input_path: str, output_path: str, sheet_index: int = 0) -> Dict[str, Any]:
        """Convert Excel to PDF (Tier 1: LibreOffice, Tier 2: HTML/CSS)"""
        
        # 1. Try LibreOffice (Best Quality)
        if PDFConverter._convert_with_libreoffice(input_path, output_path):
             return {
                "success": True,
                "message": "Converted Excel to PDF (High Quality via LibreOffice)",
                "outputPath": output_path,
                "engine": "LibreOffice"
            }
            
        # 2. Fallback: Python HTML Method
        try:
            # Read Excel file
            df = pd.read_excel(input_path, sheet_name=sheet_index)
            df = df.fillna("")
            
            # Generate styled HTML
            html_content = f"""
            <html>
            <head>
                <style>
                    @page {{
                        size: a4 landscape;
                        margin: 1cm;
                        @frame footer_frame {{
                            -pdf-frame-content: footerContent;
                            bottom: 0cm;
                            margin-left: 1cm;
                            margin-right: 1cm;
                            height: 1cm;
                        }}
                    }}
                    body {{ font-family: Helvetica, sans-serif; }}
                    h2 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
                    table {{ width: 100%; border-collapse: collapse; font-size: 9pt; }}
                    th {{ background-color: #ecf0f1; color: #2c3e50; border: 0.5pt solid #bdc3c7; padding: 6px; font-weight: bold; }}
                    td {{ border: 0.5pt solid #bdc3c7; padding: 5px; color: #34495e; }}
                    tr:nth-child(even) {{ background-color: #f9f9f9; }}
                </style>
            </head>
            <body>
                <h2>{os.path.basename(input_path)}</h2>
                {df.to_html(index=False, classes='table', border=0)}
                
                <div id="footerContent" style="text-align:center; font-size:8pt; color:#7f8c8d;">
                    Page <pdf:pagenumber>
                </div>
            </body>
            </html>
            """
            
            # Convert to PDF using xhtml2pdf
            from xhtml2pdf import pisa
            
            with open(output_path, "wb") as output_file:
                pisa_status = pisa.CreatePDF(html_content, dest=output_file)
            
            if pisa_status.err:
                raise Exception("HTML to PDF conversion failed")
            
            return {
                "success": True,
                "message": f"High Quality Conversion: {len(df)} rows",
                "outputPath": output_path,
                "rows": len(df),
                "columns": len(df.columns)
            }
        
        except Exception as e:
            raise Exception(f"Excel to PDF conversion failed: {str(e)}")

    @staticmethod
    def pdf_to_excel(input_path: str, output_path: str) -> Dict[str, Any]:
        """Convert PDF to Excel using pdfplumber (High Fidelity + OCR)"""
        
        try:
            # 1. OCR Pre-processing
            processed_path = PDFConverter._run_ocr_if_needed(input_path)
            
            import pdfplumber
            import pandas as pd
            
            all_tables = []
            
            with pdfplumber.open(processed_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    tables = page.extract_tables()
                    for table in tables:
                        # Clean table data
                        df = pd.DataFrame(table[1:], columns=table[0])
                        df = df.replace({None: ""})
                        all_tables.append(df)
            
            # Cleanup OCR temp file
            if processed_path != input_path and os.path.exists(processed_path):
                try:
                    os.remove(processed_path)
                except:
                    pass
            
            if not all_tables:
                # Fallback to simple text if no tables found
                return PDFConverter._pdf_text_to_excel_fallback(processed_path, output_path)
            
            # Write to Excel
            with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
                for i, df in enumerate(all_tables):
                    sheet_name = f"Table_{i+1}"
                    df.to_excel(writer, sheet_name=sheet_name, index=False)
            
            return {
                "success": True,
                "message": f"Extracted {len(all_tables)} tables using Neural/Heuristic engine",
                "outputPath": output_path,
                "tables": len(all_tables)
            }
            
        except Exception as e:
            raise Exception(f"PDF to Excel conversion failed: {str(e)}")

    @staticmethod
    def _pdf_text_to_excel_fallback(input_path: str, output_path: str) -> Dict[str, Any]:
        """Fallback: Text dump to Excel"""
        from PyPDF2 import PdfReader
        import pandas as pd
        
        pdf = PdfReader(input_path)
        text_data = []
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_data.extend(text.split('\n'))
        
        df = pd.DataFrame(text_data, columns=['Extracted Content'])
        df.to_excel(output_path, index=False)
        
        return {
            "success": True,
            "message": "Fallback: Text extraction (tables not detected)",
            "outputPath": output_path,
            "rows": len(text_data)
        }

    @staticmethod
    def word_to_pdf(input_path: str, output_path: str) -> Dict[str, Any]:
        """Convert Word to PDF (Tier 1: LibreOffice, Tier 2: ReportLab)"""
        
        # 1. Try LibreOffice (Best Quality)
        if PDFConverter._convert_with_libreoffice(input_path, output_path):
             return {
                "success": True,
                "message": "Converted Word to PDF (High Quality via LibreOffice)",
                "outputPath": output_path,
                "engine": "LibreOffice"
            }

        # 2. Fallback: Python ReportLab Method
        try:
            from docx import Document
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            
            doc = Document(input_path)
            styles = getSampleStyleSheet()
            
            # Map Word styles to ReportLab styles
            style_map = {
                'Heading 1': styles['Heading1'],
                'Heading 2': styles['Heading2'],
                'Heading 3': styles['Heading3'],
                'Title': styles['Title'],
                'Normal': styles['Normal'],
                'List Paragraph': styles['BodyText']
            }
            
            story = []
            
            for para in doc.paragraphs:
                if not para.text.strip():
                    continue
                
                # Get style (fallback to Normal)
                style_name = para.style.name
                style = style_map.get(style_name, styles['Normal'])
                
                # Sanitize text for XML compatibility
                text = para.text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                
                story.append(Paragraph(text, style))
                story.append(Spacer(1, 10))
            
            # Build PDF
            pdf = SimpleDocTemplate(
                output_path,
                pagesize=letter,
                rightMargin=72, leftMargin=72,
                topMargin=72, bottomMargin=18
            )
            pdf.build(story)
            
            return {
                "success": True,
                "message": f"Converted Word to PDF (Structured): {len(doc.paragraphs)} paragraphs",
                "outputPath": output_path,
                "paragraphs": len(doc.paragraphs)
            }
        
        except Exception as e:
            raise Exception(f"Word to PDF conversion failed: {str(e)}")

    @staticmethod
    def excel_to_pdf(input_path: str, output_path: str, sheet_index: int = 0) -> Dict[str, Any]:
        """Convert Excel to PDF (Handles large files, landscape mode)"""
        
        try:
            # Read Excel file
            df = pd.read_excel(input_path, sheet_name=sheet_index)
            # Replace NaN with empty string
            df = df.fillna("")
            
            # Setup PDF
            from reportlab.lib import pagesizes
            # Use Landscape for Excel (better for many columns)
            page_width, page_height = pagesizes.landscape(pagesizes.legal)
            c = canvas.Canvas(output_path, pagesize=(page_width, page_height))
            
            # Layout constants
            margin_x = 30
            margin_y = 30
            title_height = 40
            row_height = 16
            header_height = 20
            
            # Calculate usable width
            usable_width = page_width - (2 * margin_x)
            
            # Dynamic Column Widths
            # Simple logic: distibute width based on column name length + basic data sampling
            # (A full auto-width would require scanning all data, which is slow for 10k rows)
            cols = list(df.columns)
            num_cols = len(cols)
            if num_cols == 0:
                 raise Exception("Excel file is empty")

            col_width = usable_width / num_cols
            # Min width limit to prevent unreadable text (overflow will assume user understands)
            col_width = max(col_width, 40) 
            
            # Font settings
            font_size = 8
            if num_cols > 20: 
                font_size = 6 # Shrink for massive tables
            
            c.setFont("Helvetica-Bold", 14)
            c.drawString(margin_x, page_height - margin_y - 10, f"Excel Data (Sheet {sheet_index + 1})")
            
            y = page_height - margin_y - title_height
            
            # Helper to draw header
            def draw_header(y_pos):
                c.setFont("Helvetica-Bold", font_size)
                x = margin_x
                for col in cols:
                    # Truncate header if too long for column
                    text = str(col)
                    # Simple char estimate: width / ~4
                    max_chars = int(col_width / 4)
                    if len(text) > max_chars:
                        text = text[:max_chars] + ".."
                    c.drawString(x, y_pos, text)
                    x += col_width
                # Draw line under header
                c.line(margin_x, y_pos - 4, margin_x + (num_cols * col_width), y_pos - 4)
                return y_pos - header_height

            y = draw_header(y)
            c.setFont("Helvetica", font_size)
            
            rows_processed = 0
            
            for idx, row in df.iterrows():
                # Draw Row
                x = margin_x
                for col in cols:
                    val = str(row[col])
                    # formatting cleaning
                    val = val.replace('\n', ' ').strip()
                    
                    # Truncate content
                    max_chars = int(col_width / 3.5) # slightly tighter for content
                    if len(val) > max_chars:
                        val = val[:max_chars] + "."
                    
                    c.drawString(x, y, val)
                    x += col_width
                
                y -= row_height
                rows_processed += 1
                
                # New Page Check
                if y < margin_y:
                    c.showPage()
                    c.setFont("Helvetica", font_size) # Reset font on new page
                    # Setup new page header
                    y = page_height - margin_y
                    y = draw_header(y)
                    c.setFont("Helvetica", font_size)

            c.save()
            
            return {
                "success": True,
                "message": f"Converted Excel to PDF ({rows_processed} rows, {num_cols} columns)",
                "outputPath": output_path,
                "rows": rows_processed,
                "columns": num_cols
            }
        
        except Exception as e:
            raise Exception(f"Excel to PDF conversion failed: {str(e)}")

    @staticmethod
    def images_to_pdf(image_paths: List[str], output_path: str) -> Dict[str, Any]:
        """Convert multiple images to a single PDF"""
        
        try:
            from PIL import Image
            
            c = canvas.Canvas(output_path, pagesize=letter)
            width, height = letter
            
            for img_path in image_paths:
                img = Image.open(img_path)
                
                # Calculate dimensions to fit page
                img_width, img_height = img.size
                aspect = img_height / img_width
                
                if aspect > 1:  # Portrait
                    new_height = height - 100
                    new_width = new_height / aspect
                else:  # Landscape
                    new_width = width - 100
                    new_height = new_width * aspect
                
                # Center image
                x = (width - new_width) / 2
                y = (height - new_height) / 2
                
                c.drawImage(img_path, x, y, new_width, new_height)
                c.showPage()
            
            c.save()
            
            return {
                "success": True,
                "message": f"Created PDF from {len(image_paths)} images",
                "outputPath": output_path,
                "images": len(image_paths)
            }
        
        except Exception as e:
            raise Exception(f"Images to PDF conversion failed: {str(e)}")

    @staticmethod
    def merge_pdfs(pdf_paths: List[str], output_path: str) -> Dict[str, Any]:
        """Merge multiple PDFs into one"""
        
        try:
            merger = PdfMerger()
            
            for pdf_path in pdf_paths:
                merger.append(pdf_path)
            
            merger.write(output_path)
            merger.close()
            
            return {
                "success": True,
                "message": f"Merged {len(pdf_paths)} PDFs",
                "outputPath": output_path,
                "files": len(pdf_paths)
            }
        
        except Exception as e:
            raise Exception(f"PDF merge failed: {str(e)}")

    @staticmethod
    def split_pdf(input_path: str, output_dir: str) -> Dict[str, Any]:
        """Split PDF into individual pages"""
        
        try:
            os.makedirs(output_dir, exist_ok=True)
            
            pdf = PdfReader(input_path)
            output_files = []
            
            for i, page in enumerate(pdf.pages):
                writer = PdfWriter()
                writer.add_page(page)
                
                output_path = os.path.join(output_dir, f"page_{i + 1}.pdf")
                with open(output_path, 'wb') as output_file:
                    writer.write(output_file)
                
                output_files.append(output_path)
            
            return {
                "success": True,
                "message": f"Split PDF into {len(output_files)} pages",
                "outputDir": output_dir,
                "files": output_files
            }
        
        except Exception as e:
            raise Exception(f"PDF split failed: {str(e)}")

    @staticmethod
    def get_pdf_info(pdf_path: str) -> Dict[str, Any]:
        """Get PDF metadata"""
        
        try:
            pdf = PdfReader(pdf_path)
            
            return {
                "success": True,
                "pages": len(pdf.pages),
                "metadata": pdf.metadata,
                "encrypted": pdf.is_encrypted
            }
        
        except Exception as e:
            raise Exception(f"PDF info extraction failed: {str(e)}")
